# -*- coding: utf-8 -*-
"""
生成考研智能伴学系统数据库逻辑结构设计 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ==================== 页面设置 ====================
from docx.shared import Inches, Twips
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

# ==================== 字体设置 ====================
def set_run_font(run, font_name_cn='宋体', font_name_en='Times New Roman', font_size=12, bold=False):
    run.font.name = font_name_en
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)

# ==================== 三线表样式 ====================
def set_table_style(table):
    """设置三线表样式"""
    tbl = table._tbl
    tblPr = tbl.tblPr

    # 表格整体居中
    tblPr.tblJc = OxmlElement('w:tblJc')
    tblPr.tblJc.set(qn('w:val'), 'center')

    # 设置表格宽度为自适应
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)

    # 表格外边框
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        if border_name in ['top', 'bottom']:
            border.set(qn('w:sz'), '12')  # 粗线
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
        else:
            border.set(qn('w:sz'), '4')   # 细线
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = 'w:{}'.format(edge)
        element = OxmlElement(tag)
        element.set(qn('w:val'), kwargs.get(edge, 'single'))
        element.set(qn('w:sz'), kwargs.get('sz', '4'))
        element.set(qn('w:color'), kwargs.get('color', '000000'))
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_vertical_alignment(cell, alignment='center'):
    """设置单元格垂直居中"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), alignment)
    tcPr.append(vAlign)

def add_table_title(doc, cn_title, en_title):
    """添加表格标题（中文+英文）"""
    # 中文标题
    p_cn = doc.add_paragraph()
    p_cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cn.paragraph_format.space_before = Pt(6)
    p_cn.paragraph_format.space_after = Pt(0)
    run_cn = p_cn.add_run(cn_title)
    set_run_font(run_cn, font_size=12, bold=True)

    # 英文标题
    p_en = doc.add_paragraph()
    p_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_en.paragraph_format.space_before = Pt(0)
    p_en.paragraph_format.space_after = Pt(6)
    run_en = p_en.add_run(en_title)
    set_run_font(run_en, font_size=12, bold=False)

def add_table_intro(doc, text):
    """添加表格说明段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, font_size=12, bold=False)
    return p

def add_page_break_before(doc):
    """在段落前添加分页"""
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    return p

def create_table_header(table, headers, col_widths):
    """创建表头行"""
    header_cells = table.rows[0].cells
    for i, (cell, header, width) in enumerate(zip(header_cells, headers, col_widths)):
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_vertical_alignment(cell, 'center')

        run = cell.paragraphs[0].runs[0]
        set_run_font(run, font_size=10.5, bold=True)

        # 设置单元格宽度
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(width * 567)))  # cm -> twips
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

        # 表头下边框
        set_cell_border(cell, top='single', sz='8', color='000000',
                        left='single', right='single', bottom='single')

def add_data_row(table, row_idx, data, col_widths):
    """添加数据行"""
    row = table.rows[row_idx]
    for i, (cell, value, width) in enumerate(zip(row.cells, data, col_widths)):
        cell.text = str(value) if value else ''
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_vertical_alignment(cell, 'center')

        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run('')
        set_run_font(run, font_size=10.5, bold=False)

        # 设置单元格宽度
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(width * 567)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

        # 设置浅色内部边框
        set_cell_border(cell, top='single', sz='4', color='000000',
                        left='single', right='single', bottom='single')

# ==================== 表格定义数据 ====================

# 列宽比例（总约14.66cm，用于正文区）
COL_WIDTHS = [2.1, 1.8, 0.9, 1.1, 1.2, 2.8, 2.76]  # 字段名、数据类型、长度、允许为空、键引用、字段含义、备注

TABLES = [
    {
        "cn": "表4-1  用户信息表",
        "en": "Tab. 4-1 User Information Table",
        "intro": "user数据表主要用于记录系统考研用户的基本身份信息与账号状态，包括用户编号、用户名、密码、姓名、头像、角色、电话、邮箱和账号状态等属性，用于实现用户登录认证、身份区分及个人信息管理等功能，具体如表4-1所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "INT", "11", "否", "主键", "用户编号", "自增"],
            ["username", "VARCHAR", "255", "是", "", "账号", "唯一"],
            ["password", "VARCHAR", "255", "是", "", "密码", "BCrypt加密存储"],
            ["name", "VARCHAR", "255", "是", "", "姓名", ""],
            ["avatar", "VARCHAR", "255", "是", "", "头像", "存储图片URL"],
            ["role", "VARCHAR", "255", "是", "", "角色", "用户角色标识"],
            ["status", "INT", "11", "是", "", "账号状态", "1-正常，0-禁用"],
            ["phone", "VARCHAR", "255", "是", "", "电话", ""],
            ["email", "VARCHAR", "255", "是", "", "邮箱", ""],
        ]
    },
    {
        "cn": "表4-2  管理员信息表",
        "en": "Tab. 4-2 Administrator Information Table",
        "intro": "admin数据表主要用于存储系统管理员的账号信息与权限相关属性，包括管理员编号、账号、密码、姓名、头像、角色、电话、邮箱等字段，用于后台系统登录认证与权限管理，具体如表4-2所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "INT", "11", "否", "主键", "管理员编号", "自增"],
            ["username", "VARCHAR", "255", "是", "", "账号", "唯一"],
            ["password", "VARCHAR", "255", "是", "", "密码", "BCrypt加密存储"],
            ["name", "VARCHAR", "255", "是", "", "姓名", ""],
            ["avatar", "VARCHAR", "255", "是", "", "头像", "存储图片URL"],
            ["role", "VARCHAR", "255", "是", "", "角色", "管理员角色标识"],
            ["phone", "VARCHAR", "255", "是", "", "电话", ""],
            ["email", "VARCHAR", "255", "是", "", "邮箱", ""],
            ["status", "INT", "11", "是", "", "账号状态", "1-正常，0-禁用"],
        ]
    },
    {
        "cn": "表4-3  院校信息表",
        "en": "Tab. 4-3 University Information Table",
        "intro": "university数据表用于存储全国各高等院校的详细信息，是系统的核心数据表之一。包括院校编号、名称、logo、所在地区、详细地址、院校类型、办学层次、是否985/211/双一流、官网地址、院校简介等基本信息，以及学校介绍、校园环境、图书馆、宿舍、食堂、交通等学习生活相关信息，还包含硕士培养信息、重点学科、特色专业、一流学科、考研热门专业等考研相关字段，实现院校信息的完整管理和展示，具体如表4-3所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "INT", "11", "否", "主键", "院校编号", "自增"],
            ["name", "VARCHAR", "255", "是", "", "院校名称", ""],
            ["avatar", "VARCHAR", "255", "是", "", "校徽/头像", "存储图片URL"],
            ["province_id", "INT", "11", "是", "外键(areas.id)", "所在省份ID", "关联地区表"],
            ["address", "VARCHAR", "255", "是", "", "详细地址", ""],
            ["school_type", "VARCHAR", "50", "是", "", "院校类型", "综合类/理工类/师范类等"],
            ["education_level", "VARCHAR", "50", "是", "", "办学层次", "本科/专科"],
            ["is_985", "TINYINT", "1", "否", "", "是否985", "0-否，1-是"],
            ["is_211", "TINYINT", "1", "否", "", "是否211", "0-否，1-是"],
            ["is_double_first_class", "TINYINT", "1", "否", "", "是否双一流", "0-否，1-是"],
            ["official_website", "VARCHAR", "255", "是", "", "官方网址", "以http(s)://开头"],
            ["description", "LONGTEXT", "", "是", "", "院校简介", ""],
            ["longitude", "DECIMAL", "10,6", "是", "", "百度经度", "BD09坐标系"],
            ["latitude", "DECIMAL", "10,6", "是", "", "百度纬度", "BD09坐标系"],
            ["english_name", "VARCHAR", "255", "是", "", "英文名称", ""],
            ["founded_year", "VARCHAR", "50", "是", "", "建校年份", ""],
            ["affiliation", "VARCHAR", "255", "是", "", "隶属单位", ""],
            ["leader_info", "TEXT", "", "是", "", "领导信息", ""],
            ["school_intro", "TEXT", "", "是", "", "学校介绍", ""],
            ["campus_environment", "TEXT", "", "是", "", "校园环境", ""],
            ["contact_address", "VARCHAR", "500", "是", "", "通讯地址", ""],
            ["contact_phone", "VARCHAR", "50", "是", "", "联系电话", ""],
            ["contact_email", "VARCHAR", "100", "是", "", "联系邮箱", ""],
            ["postcode", "VARCHAR", "20", "是", "", "邮政编码", ""],
            ["library_info", "TEXT", "", "是", "", "图书馆信息", ""],
            ["dormitory_info", "TEXT", "", "是", "", "宿舍信息", ""],
            ["canteen_info", "TEXT", "", "是", "", "食堂信息", ""],
            ["transport_info", "TEXT", "", "是", "", "交通信息", ""],
            ["master_program_info", "TEXT", "", "是", "", "硕士培养信息", ""],
            ["key_discipline_info", "TEXT", "", "是", "", "重点学科信息", ""],
            ["featured_major_info", "TEXT", "", "是", "", "特色专业信息", ""],
            ["first_class_major_info", "TEXT", "", "是", "", "一流学科信息", ""],
            ["graduate_hot_major_info", "TEXT", "", "是", "", "考研热门专业信息", ""],
            ["create_time", "DATETIME", "", "否", "", "创建时间", "自动创建"],
            ["update_time", "DATETIME", "", "否", "", "更新时间", "自动更新"],
        ]
    },
    {
        "cn": "表4-4  地区信息表",
        "en": "Tab. 4-4 Area Information Table",
        "intro": "areas数据表用于存储全国各省份及直辖市等地区信息，属于系统的基础数据表。通过province_id字段与university等表关联，实现按地区筛选院校的功能，具体如表4-4所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "INT", "11", "否", "主键", "地区编号", "自增"],
            ["name", "VARCHAR", "255", "是", "", "地区名称", "如：北京市、山东省"],
        ]
    },
    {
        "cn": "表4-5  学科门类表",
        "en": "Tab. 4-5 Academic Discipline Table",
        "intro": "categorys数据表用于存储教育部规定的学科门类信息，如哲学、经济学、法学、教育学、文学、历史学、理学、工学、农学、医学、军事学、管理学、艺术学等。是专业信息表specialtys的上级字典表，通过categorys_id字段建立层级关系，具体如表4-5所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "INT", "11", "否", "主键", "门类编号", "自增"],
            ["name", "VARCHAR", "255", "是", "", "门类名称", "如：工学、理学"],
        ]
    },
    {
        "cn": "表4-6  专业信息表",
        "en": "Tab. 4-6 Specialty Information Table",
        "intro": "specialtys数据表用于存储各学科门类下的具体专业信息，包括专业编号、所属门类ID、专业名称、专业代码、专业简介、培养目标、主要课程、就业方向及官方来源链接等字段。通过categorys_id关联学科门类表，通过university_specialtys中间表关联院校，实现专业信息的完整管理，具体如表4-6所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "INT", "11", "否", "主键", "专业编号", "自增"],
            ["categorys_id", "INT", "11", "是", "外键(categorys.id)", "所属门类ID", "关联学科门类表"],
            ["name", "VARCHAR", "255", "是", "", "专业名称", ""],
            ["specialty_number", "VARCHAR", "255", "是", "", "专业代码", "如：081200"],
            ["intro", "TEXT", "", "是", "", "专业简介", ""],
            ["training_objective", "TEXT", "", "是", "", "培养目标", ""],
            ["main_courses", "TEXT", "", "是", "", "主要课程", ""],
            ["employment_direction", "TEXT", "", "是", "", "就业方向", ""],
            ["official_source_url", "VARCHAR", "255", "是", "", "官方来源链接", ""],
        ]
    },
    {
        "cn": "表4-7  院校专业关联表",
        "en": "Tab. 4-7 University-Specialty Relation Table",
        "intro": "university_specialtys数据表是院校与专业之间的多对多关联表。每条记录表示某院校开设的某个专业，包含专业在该院校的具体介绍内容content字段。通过university_id关联院校表，通过categorys_id关联门类表，通过specialtys_id关联专业信息表，实现院校专业的灵活配置，具体如表4-7所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "INT", "11", "否", "主键", "编号", "自增"],
            ["university_id", "INT", "11", "是", "外键(university.id)", "院校ID", "关联院校信息表"],
            ["categorys_id", "INT", "11", "是", "外键(categorys.id)", "门类ID", "关联学科门类表"],
            ["specialtys_id", "INT", "11", "是", "外键(specialtys.id)", "专业ID", "关联专业信息表"],
            ["content", "LONGTEXT", "", "是", "", "专业介绍", "该专业在指定院校的介绍内容"],
        ]
    },
    {
        "cn": "表4-8  招生政策表",
        "en": "Tab. 4-8 Admission Policy Table",
        "intro": "policys数据表用于存储各高等院校的招生政策信息，包括政策编号、标题、关联院校ID、政策简介、详细内容、浏览量及发布时间等字段。考生可通过招生政策模块了解各院校的招生要求、考试科目、分数线等关键信息，具体如表4-8所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "INT", "11", "否", "主键", "政策编号", "自增"],
            ["name", "VARCHAR", "255", "是", "", "政策标题", ""],
            ["university_id", "INT", "11", "是", "外键(university.id)", "院校ID", "关联院校信息表"],
            ["intro", "VARCHAR", "255", "是", "", "政策简介", ""],
            ["content", "LONGTEXT", "", "是", "", "政策详情", "富文本内容"],
            ["view_count", "INT", "11", "是", "", "浏览量", "点击次数"],
            ["time", "VARCHAR", "255", "是", "", "创建时间", ""],
        ]
    },
    {
        "cn": "表4-9  院校评价表",
        "en": "Tab. 4-9 University Comment Table",
        "intro": "comment数据表用于存储注册用户对各高等院校的评价信息，包括评价编号、关联院校ID、评价用户ID、评价内容、评分及评价时间等字段。考生可通过查看其他用户的评价了解院校的真实就读体验，具体如表4-9所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "INT", "11", "否", "主键", "评价编号", "自增"],
            ["university_id", "INT", "11", "是", "外键(university.id)", "院校ID", "关联院校信息表"],
            ["user_id", "INT", "11", "是", "外键(user.id)", "用户ID", "关联用户信息表"],
            ["details", "VARCHAR", "255", "是", "", "评价内容", ""],
            ["mark", "INT", "11", "是", "", "评分", "1-5分制"],
            ["time", "VARCHAR", "255", "是", "", "评价时间", ""],
        ]
    },
    {
        "cn": "表4-10  专业解读表",
        "en": "Tab. 4-10 Major Interpretation Table",
        "intro": "interpretations数据表用于存储各专业方向的深度解读文章，包括解读编号、标题、所属门类ID、简介、封面图片、详细内容、浏览量及发布时间等字段。考生可通过专业解读模块深入了解各专业的培养方案、研究方向及就业前景，具体如表4-10所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "INT", "11", "否", "主键", "解读编号", "自增"],
            ["name", "VARCHAR", "255", "是", "", "解读标题", ""],
            ["categorys_id", "INT", "11", "是", "外键(categorys.id)", "门类ID", "关联学科门类表"],
            ["intro", "VARCHAR", "2000", "是", "", "解读简介", ""],
            ["img", "VARCHAR", "255", "是", "", "封面图片", "存储图片URL"],
            ["content", "LONGTEXT", "", "是", "", "解读详情", "富文本内容"],
            ["view_count", "INT", "11", "是", "", "浏览量", "点击次数"],
            ["time", "VARCHAR", "255", "是", "", "发布时间", ""],
        ]
    },
    {
        "cn": "表4-11  学习计划主表",
        "en": "Tab. 4-11 Study Plan Main Table",
        "intro": "study_plan数据表是学习计划功能的核心主表，用于记录用户每日学习计划的整体信息。包括计划编号、关联用户ID、计划日期、用户反馈、AI建议、计划状态、创建时间和更新时间等字段。每条学习计划记录对应一天的复习安排，具体如表4-11所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "BIGINT", "20", "否", "主键", "计划编号", "自增"],
            ["user_id", "BIGINT", "20", "否", "外键(user.id)", "用户ID", "关联用户信息表"],
            ["plan_date", "DATE", "", "否", "", "计划日期", "对应学习日期"],
            ["user_feedback", "TEXT", "", "是", "", "用户反馈", "昨日总结/当前状态"],
            ["ai_advice", "TEXT", "", "是", "", "AI建议", "AI根据反馈生成的建议"],
            ["plan_status", "VARCHAR", "20", "否", "", "计划状态", "PENDING-待生成，GENERATED-已生成"],
            ["create_time", "DATETIME", "", "否", "", "创建时间", "自动创建"],
            ["update_time", "DATETIME", "", "否", "", "更新时间", "自动更新"],
        ]
    },
    {
        "cn": "表4-12  学习任务明细表",
        "en": "Tab. 4-12 Study Plan Task Detail Table",
        "intro": "study_plan_task数据表用于存储学习计划中每日的具体学习任务明细，是study_plan的子表。每条记录表示一个具体的学习任务，包含所属计划ID、科目、任务内容、是否完成、排序号及任务来源等字段。通过task_id实现任务级别的唯一标识，支持任务的后延管理，具体如表4-12所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "BIGINT", "20", "否", "主键", "任务编号", "自增"],
            ["task_id", "VARCHAR", "64", "否", "唯一键", "任务ID", "业务唯一标识"],
            ["plan_id", "BIGINT", "20", "否", "外键(study_plan.id)", "计划ID", "关联学习计划主表，级联删除"],
            ["subject", "VARCHAR", "50", "否", "", "科目", "如：英语、数学、政治"],
            ["content", "TEXT", "", "否", "", "任务内容", "具体学习任务描述"],
            ["completed", "TINYINT", "1", "否", "", "是否完成", "0-未完成，1-已完成"],
            ["sort_no", "INT", "11", "否", "", "排序号", "同计划内任务排序"],
            ["task_source", "VARCHAR", "20", "否", "", "任务来源", "GENERATED-生成计划，DEFERRED-后延"],
            ["create_time", "DATETIME", "", "否", "", "创建时间", "自动创建"],
            ["update_time", "DATETIME", "", "否", "", "更新时间", "自动更新"],
        ]
    },
    {
        "cn": "表4-13  题目信息表",
        "en": "Tab. 4-13 Question Information Table",
        "intro": "question数据表用于存储考研各科目的练习题目，是每日一题和在线考试功能的核心数据表。包括题目编号、科目、题型、题干内容、各选项、正确答案、题目解析、难度等级、状态及时间戳等字段，支持单选、判断、简答等多种题型，具体如表4-13所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "BIGINT", "20", "否", "主键", "题目编号", "自增"],
            ["subject", "VARCHAR", "50", "否", "", "科目", "英语/数学/政治/专业课"],
            ["type", "VARCHAR", "30", "否", "", "题型", "single_choice/judge/short_answer"],
            ["title", "TEXT", "", "否", "", "题目内容", "题干文字"],
            ["option_a", "VARCHAR", "500", "是", "", "选项A", ""],
            ["option_b", "VARCHAR", "500", "是", "", "选项B", ""],
            ["option_c", "VARCHAR", "500", "是", "", "选项C", ""],
            ["option_d", "VARCHAR", "500", "是", "", "选项D", ""],
            ["correct_answer", "VARCHAR", "100", "否", "", "正确答案", ""],
            ["analysis", "TEXT", "", "是", "", "题目解析", ""],
            ["difficulty", "VARCHAR", "20", "是", "", "难度等级", "easy/medium/hard"],
            ["status", "TINYINT", "1", "是", "", "状态", "1-启用，0-禁用"],
            ["create_time", "DATETIME", "", "是", "", "创建时间", "自动创建"],
            ["update_time", "DATETIME", "", "是", "", "更新时间", "自动更新"],
        ]
    },
    {
        "cn": "表4-14  每日一题记录表",
        "en": "Tab. 4-14 Daily Question Record Table",
        "intro": "question_record数据表用于记录每位用户每日一题的作答历史，包括记录编号、用户ID、题目ID、用户答案、是否正确、作答时间及题目对应日期等字段。系统通过user_id和question_date的联合唯一索引确保每个用户每天只有一条作答记录，用于答题统计和历史查看，具体如表4-14所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "BIGINT", "20", "否", "主键", "记录编号", "自增"],
            ["user_id", "BIGINT", "20", "否", "外键(user.id)", "用户ID", "关联用户信息表"],
            ["question_id", "BIGINT", "20", "否", "外键(question.id)", "题目ID", "关联题目信息表"],
            ["user_answer", "VARCHAR", "100", "是", "", "用户答案", ""],
            ["is_correct", "TINYINT", "1", "是", "", "是否正确", "0-否，1-是"],
            ["answer_time", "DATETIME", "", "是", "", "作答时间", "自动记录"],
            ["question_date", "DATE", "", "否", "唯一键(联合)", "题目日期", "对应题目分发日期"],
        ]
    },
    {
        "cn": "表4-15  每日题目分发表",
        "en": "Tab. 4-15 Daily Question Distribution Table",
        "intro": "daily_question数据表用于管理系统每日向用户分发的推荐题目，包括分发编号、题目日期、题目ID及创建时间等字段。通过question_date字段的唯一索引确保每天只分发一道题目，管理员可提前配置未来日期的每日一题，实现题目的定时推送功能，具体如表4-15所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "BIGINT", "20", "否", "主键", "分发编号", "自增"],
            ["question_date", "DATE", "", "否", "唯一键", "题目日期", "每日期对应一道题"],
            ["question_id", "BIGINT", "20", "否", "外键(question.id)", "题目ID", "关联题目信息表"],
            ["create_time", "DATETIME", "", "是", "", "创建时间", "自动创建"],
        ]
    },
    {
        "cn": "表4-16  模拟考试成绩表",
        "en": "Tab. 4-16 Exam Result Table",
        "intro": "exam_result数据表用于记录用户参加在线模拟考试的考试成绩信息，包括成绩编号、用户ID、考试科目、试题出处、得分、答题时长、模拟模式及考试时间等字段。用户可在考前进行各科目的模拟练习，系统自动保存成绩供用户回顾分析，具体如表4-16所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "BIGINT", "20", "否", "主键", "成绩编号", "自增"],
            ["user_id", "BIGINT", "20", "否", "外键(user.id)", "用户ID", "关联用户信息表"],
            ["subject", "VARCHAR", "50", "否", "", "考试科目", "英语/数学/政治/专业课"],
            ["question_source", "VARCHAR", "255", "是", "", "试题出处", "试卷来源说明"],
            ["score", "INT", "11", "否", "", "得分", "考试成绩"],
            ["duration_seconds", "INT", "11", "否", "", "答题时长", "单位：秒"],
            ["simulation_mode", "VARCHAR", "50", "是", "", "模拟模式", "如：演示模式"],
            ["create_time", "DATETIME", "", "是", "", "考试时间", "自动记录"],
        ]
    },
    {
        "cn": "表4-17  知识库文件表",
        "en": "Tab. 4-17 Knowledge Base File Table",
        "intro": "kb_file数据表是知识库模块的文件管理主表，用于存储上传到系统的各类文档（如PDF、TXT、MD格式）的元信息。包括文件编号、标题、原始文件名、访问地址、存储路径、文件类型、业务类型、文件SHA-256哈希值、文件大小、处理状态、分片数量、Redis追踪前缀、错误信息、备注及创建人和时间戳等完整字段，实现知识库文档的完整生命周期管理，具体如表4-17所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "BIGINT", "20", "否", "主键", "文件编号", "自增"],
            ["title", "VARCHAR", "255", "否", "", "文档标题", ""],
            ["file_name", "VARCHAR", "255", "否", "", "原始文件名", ""],
            ["file_url", "VARCHAR", "500", "是", "", "文件访问地址", ""],
            ["file_path", "VARCHAR", "500", "是", "", "服务器存储路径", ""],
            ["file_type", "VARCHAR", "50", "是", "", "文件类型", "如：pdf/txt/md"],
            ["business_type", "VARCHAR", "100", "是", "", "业务类型", "如：school/policy/major"],
            ["file_hash", "VARCHAR", "64", "否", "唯一键", "文件SHA-256", "用于去重"],
            ["file_size", "BIGINT", "20", "是", "", "文件大小", "单位：字节"],
            ["status", "VARCHAR", "50", "否", "", "处理状态", "UPLOADING/PARSED/INDEXED/FAILED/DELETED"],
            ["segment_count", "INT", "11", "是", "", "分片数量", "切分后的片段总数"],
            ["redis_prefix", "VARCHAR", "255", "是", "", "Redis追踪前缀", ""],
            ["error_message", "TEXT", "", "是", "", "错误信息", "处理失败原因"],
            ["remark", "VARCHAR", "500", "是", "", "备注", ""],
            ["create_by", "BIGINT", "20", "是", "", "创建人", "管理员ID"],
            ["create_time", "DATETIME", "", "否", "", "创建时间", "自动创建"],
            ["update_time", "DATETIME", "", "否", "", "更新时间", "自动更新"],
        ]
    },
    {
        "cn": "表4-18  知识库分片表",
        "en": "Tab. 4-18 Knowledge Base Content Segment Table",
        "intro": "kb_content数据表用于存储知识库文档经过分片处理后的文本内容，是kb_file的子表。每条记录对应一个文件在某分片序号下的文本内容，包含文件ID、分片序号、分片文本、Redis追踪key及创建时间等字段。通过file_id和segment_no的联合唯一索引确保分片不重复，支持向量检索等AI功能，具体如表4-18所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "BIGINT", "20", "否", "主键", "分片编号", "自增"],
            ["file_id", "BIGINT", "20", "否", "外键(kb_file.id)", "文件ID", "关联知识库文件表"],
            ["segment_no", "INT", "11", "否", "唯一键(联合)", "分片序号", "同文件内从1开始"],
            ["content_text", "LONGTEXT", "", "是", "", "分片文本", ""],
            ["redis_key", "VARCHAR", "255", "是", "", "Redis追踪key", ""],
            ["create_time", "DATETIME", "", "否", "", "创建时间", "自动创建"],
        ]
    },
    {
        "cn": "表4-19  收藏信息表",
        "en": "Tab. 4-19 Collection Information Table",
        "intro": "collect数据表用于存储用户收藏院校的记录信息，包括收藏编号、关联院校ID、关联用户ID等字段。用户可在院校详情页面对感兴趣的院校进行收藏操作，便于后续快速查看和对比分析，具体如表4-19所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "INT", "11", "否", "主键", "收藏编号", "自增"],
            ["university_id", "INT", "11", "是", "外键(university.id)", "院校ID", "关联院校信息表"],
            ["user_id", "INT", "11", "是", "外键(user.id)", "用户ID", "关联用户信息表"],
        ]
    },
    {
        "cn": "表4-20  轮播图表",
        "en": "Tab. 4-20 Slideshow Table",
        "intro": "slideshow数据表用于管理系统首页的轮播图资源，包括轮播图编号、关联院校ID及图片地址等字段。管理员可在后台配置多条轮播图，每张轮播图可关联指定院校，点击后可跳转至对应院校详情页，实现首页的动态内容展示，具体如表4-20所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "INT", "11", "否", "主键", "轮播图编号", "自增"],
            ["university_id", "INT", "11", "是", "外键(university.id)", "院校ID", "关联院校信息表，可为空表示系统公告"],
            ["img", "VARCHAR", "255", "是", "", "轮播图图片", "存储图片URL"],
        ]
    },
    {
        "cn": "表4-21  公告信息表",
        "en": "Tab. 4-21 Notice Information Table",
        "intro": "notice数据表用于存储系统公告信息，包括公告编号、标题、内容及发布时间等字段。管理员可在后台发布系统公告，用户登录后可查看最新公告内容，实现系统重要信息的推送与展示，具体如表4-21所示。",
        "cols": ["字段名", "数据类型", "长度", "允许为空", "键引用", "字段含义", "备注"],
        "rows": [
            ["id", "INT", "11", "否", "主键", "公告编号", "自增"],
            ["title", "VARCHAR", "255", "是", "", "公告标题", ""],
            ["content", "TEXT", "", "是", "", "公告内容", ""],
            ["time", "VARCHAR", "255", "是", "", "发布时间", ""],
        ]
    },
]

# ==================== 生成文档 ====================

# 章节标题
chapter_title = doc.add_paragraph()
chapter_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
chapter_title.paragraph_format.space_before = Pt(0)
chapter_title.paragraph_format.space_after = Pt(12)
run = chapter_title.add_run("第4章  数据库逻辑结构设计")
set_run_font(run, font_size=16, bold=True)

# 章节引言
intro_text = """数据库是考研智能伴学系统的核心组成部分，其逻辑结构设计直接影响系统的数据存储效率、查询性能及扩展能力。本系统采用MySQL 8.0作为关系型数据库，基于业务功能模块划分数据表，共计21张数据表，涵盖用户管理、院校信息、专业管理、学习计划、题库练习、知识库及系统配置等核心业务领域。各数据表通过主键和外键约束建立关联关系，确保数据的完整性和一致性。本章将详细描述各数据表的字段设计、数据类型及业务含义。"""
intro_p = doc.add_paragraph()
intro_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro_p.paragraph_format.space_after = Pt(12)
run = intro_p.add_run(intro_text)
set_run_font(run, font_size=12, bold=False)

for i, table_data in enumerate(TABLES):
    # 分页符（每张表前）
    p_br = doc.add_paragraph()
    p_br.paragraph_format.page_break_before = True

    # 表标题
    add_table_title(doc, table_data["cn"], table_data["en"])

    # 说明段落
    add_table_intro(doc, table_data["intro"])

    # 创建表格
    num_rows = len(table_data["rows"]) + 1  # +1 表头
    num_cols = len(table_data["cols"])
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格样式
    set_table_style(table)

    # 表头行
    create_table_header(table, table_data["cols"], COL_WIDTHS)

    # 数据行
    for row_idx, row_data in enumerate(table_data["rows"]):
        add_data_row(table, row_idx + 1, row_data, COL_WIDTHS)

    # 设置表格行不允许跨页断行
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), '1')
        trPr.append(cantSplit)

    # 表头重复（对跨页表格有效）
    tbl = table._tbl
    tblHeader = OxmlElement('w:tblHeader')
    tblPr = tbl.tblPr
    tblPr.append(tblHeader)

    # 段后间距
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(6)

# 保存文档
output_path = r"c:\Users\Lenovo\Desktop\kaoyan_guide\数据库逻辑结构设计_论文表格版.docx"
doc.save(output_path)
print(f"文档已生成：{output_path}")
print(f"共计 {len(TABLES)} 张数据表")
